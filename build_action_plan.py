#!/usr/bin/env python3
"""Build the unified Action Plan + Ready-to-Ship Copy doc as a .docx."""
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11)
ACCENT = RGBColor(0x1A, 0x53, 0x4F); MUTED = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xB0, 0x00, 0x00); GREEN = RGBColor(0x1E, 0x7A, 0x33)

def title(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(24)
def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)
def body(t): return doc.add_paragraph(t)
def sub(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = True; r.font.color.rgb = MUTED; return p
def ship(t):
    """Ready-to-ship copy block, boxed look via bold accent + indent."""
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
    r = p.add_run(t); r.font.size = Pt(12); return p
def label(lbl, value):
    p = doc.add_paragraph(); r = p.add_run(f"{lbl}: "); r.bold = True; r.font.color.rgb = ACCENT
    p.add_run(value); return p
def bullets(items):
    for it in items: doc.add_paragraph(it, style="List Bullet")
def numbered(items):
    for it in items: doc.add_paragraph(it, style="List Number")
def flag(t):
    p = doc.add_paragraph(); r = p.add_run(f"⚑ {t}"); r.bold = True; r.font.color.rgb = RED
def note(t):
    p = doc.add_paragraph(); r = p.add_run(f"→ {t}"); r.italic = True; r.font.color.rgb = MUTED

# =====================================================================
title("Well and Good Websites: Action Plan & Ready-to-Ship Copy")
sub("Built from your live-site audit + competitive teardown (Cedar+Mint, KMK, Cool Koala), 2026-06-28. "
    "Strategy in one line: you have the strongest offer and the weakest credibility, so this is mostly about "
    "making a strong offer believable, not repositioning.")

h2("The sequence that matters")
body("Nothing below the credibility line converts until the credibility line is done. Order is deliberate.")
numbered([
    "Relabel / move the github.io prototypes, they currently read as fake client work (live liability).",
    "Add the founder block (About Matt + photo + Jetta Grove track record).",
    "Add Frank's case study + approved testimonial, your one real, live client.",
    "Launch the guarantee, open ground none of your four competitors hold.",
    "Fix GA4 / Hotjar, you're optimizing blind until this works.",
    ", credibility line, everything below only pays off once the above is true, ",
    "Fix the H1, ship the new hero, add the proof bar, add booking, start the blog.",
])

# =====================================================================
h1("Part 1, Ready-to-ship copy")
sub("Write-ready. Bracketed items need your one input or Frank's approval. Nothing fabricated.")

# --- Hero ---
h2("Hero headline, 5 variants to test")
sub("Framework rule: action verb + specific outcome + timeframe/contrast. Your current hero leads with a product "
    "attribute; these lead with the customer outcome and fold in your real wedges (24h, bundle, found-on-Google, price).")
body("A, Outcome + timeframe (recommended primary):")
ship("Get found on Google and chosen by local customers, your site live in 24 hours.")
body("B, Price contrast:")
ship("The website an agency would charge $5,000 for. Live in 24 hours, from $99/month.")
body("C, Risk-reversal led (recommended challenger):")
ship("See your new website before you pay a cent. Live in 24 hours, from $99/month.")
body("D, Bundle:")
ship("Website, SEO, and social, for less than agencies charge for SEO alone. Live in 24 hours.")
body("E, Pain / enemy:")
ship("Stop losing local searches to the competitor who showed up first.")
note("Test A vs C first (outcome vs risk-reversal). Keep the H2 confession line, it's already excellent: "
     "\"Local agencies don't publish their pricing. We have nothing to hide.\"")

h2("Eyebrow / section line, replace the filler")
body("Current \"Get more attention. Get more customers.\" says nothing only you can say. Replace with:")
ship("Everything an agency sells separately, website, SEO, social, reviews, bundled into one monthly price.")

# --- Guarantee ---
h2("Guarantee (you confirmed both)")
body("Primary line (hero sub-CTA + repeat at the pricing table):")
ship("Your site live in 24 hours, or your first month is free. Then I revise it until you're happy.")
note("This is the only risk-reversal in your competitive set. Only ship the 24-hour half if you can reliably hit it; "
     "the 'revise until happy' half is safe regardless.")

# --- About Matt ---
h2("About Matt, founder block")
flag("[INSERT HEADSHOT, the photo you sent. Place left of the text, ~150px round.]")
body("Headline:")
ship("Hi, I'm Matt. I think your neighbourhood deserves the same tech the big players use.")
body("Body:")
ship("I'm Matthew Baggetta. Through my consultancy, Jetta Grove, I've spent years putting frontier AI and the "
     "latest growth techniques to work for companies that can afford them.")
ship("Here's what I kept noticing: the big budgets and the high-tech industries pull further ahead every month, "
     "because they're the only ones using this stuff. So I started Well and Good Websites to hand that same fire to "
     "the businesses on my own street.")
ship("I live in Welland. I want to see it do well. So I use the exact techniques and frontier models the big players "
     "guard, and I point them at one thing: getting local shops, trades, and restaurants found, chosen, and growing.")
ship("Same tech as the giants. Built for the business down the block.")
note("Optional bolder pull-quote if it fits the brand: \"My goal is simple, steal the fire from the gods and hand "
     "it to Main Street.\" Use sparingly; it's a strong line but check it matches your voice.")

h2("Founder proof bar (from Jetta Grove, honestly attributed)")
body("Place under the hero CTA or inside the About block. This fixes the 'unattributed stats' liability by tying the "
     "numbers to you, not to an implied Well and Good client:")
ship("Through my consultancy, Jetta Grove, I've driven 4.34M impressions, ranked for 163 keywords, and put work on "
     "Page 1 of Google. Now I bring that to local business.")
flag("[ADD A SCREENSHOT or link to jettagrove.com next to this, it's what makes the numbers believable.]")
note("Because Jetta Grove is your own consultancy, this is true and defensible. Confirm the work is SEO/discovery-related "
     "so it directly backs the 'get found on Google & ChatGPT' promise; if the niche is unrelated, we'll frame it as "
     "proof of skill rather than proof of local results.")

# --- Recent Work ---
h2("Recent Work, honest relabel (most urgent fix)")
flag("Two of three items live on immeasurablematt.github.io and read as spec builds, which contradicts your "
     "\"agency results\" claim. Fix this first.")
h3("Item 1, Frank Baggetta (REAL, feature first)")
ship("Frank Baggetta · live for months · FrankBaggetta.ca")
note("This is your only real, live client. Make it visually the lead, with the testimonial below. Open the live link "
     "in a NEW TAB so you don't leak the visitor off-funnel.")
h3("Items 2 & 3, Sandwich shop & auto repair (PROTOTYPES)")
body("Until they buy, label them honestly. Two options:")
bullets([
    "Best: move them to your own subdomains (e.g. demo.wellandgoodwebsites.ca/sandwich) and tag them \"Concept build.\"",
    "Minimum: keep where they are but add a clear \"Concept build, example, not a client\" label.",
])
body("Suggested labels:")
ship("Concept build, a sandwich shop in Welland. An example of what a one-page site looks like for a quick-service spot.")
ship("Concept build, an auto shop in Welland. An example of a call-first layout built around real Google reviews.")
note("If either client buys this week, promote it to a real case study with their name + a result.")

# --- Frank testimonial ---
h2("Frank's testimonial, DRAFT for his approval")
flag("[SEND TO FRANK TO APPROVE OR EDIT. Do not publish until he signs off. Swap in a specific result if he has one, "
     "more inquiries, found on Google for his name, etc.]")
ship("\"I'm not a tech person, and I'd put off a proper website for too long. Matt had mine live faster than I thought "
     "was possible, and now it actually looks the part when someone Googles me. I sent him what I had and he handled "
     "the rest, painless.\", Frank Baggetta, FrankBaggetta.ca")
note("I kept this experiential and metric-free on purpose, I won't invent numbers for Frank. If he can give one real "
     "outcome, it gets far stronger (the teardown's point: proof-with-numbers wins ties vs. Cool Koala).")

# --- H1 fix ---
h2("H1 fix (SEO)")
body("Your teardown found the semantic H1 is \"Good websites. Done well.\" (no keywords), competing with the hero. Make "
     "the keyword hero the single H1 and demote the tagline to styled text:")
ship("H1: Get found on Google and chosen by local customers, your site live in 24 hours.")
ship("Tagline (not an H1): Good websites. Done well.")

# =====================================================================
h1("Part 2, Implementation checklist")
h2("This week (credibility, do in order)")
numbered([
    "Relabel/move the two prototype sites off github.io; open all 'view live' links in a new tab.",
    "Publish the About Matt block with your photo + the Jetta Grove proof bar.",
    "Add Frank's case study; send him the draft testimonial; publish once approved.",
    "Add the guarantee line to the hero and the pricing table.",
    "Fix GA4 + Hotjar; confirm events fire before running any test.",
    "Swap the H1 per above.",
])
h2("This month (amplify)")
bullets([
    "Collect 5–10 Google reviews → embed the star badge in the hero (KMK and Cool Koala both lead with 5.0★).",
    "Turn Frank, and any prototype that closes, into a numbers case study (situation → build → result + screenshot).",
    "Add a Cal.com 'Book a 15-min call' button beside the form (Cool Koala already books instantly; you make leads wait).",
    "Publish 3 blog posts to prove the content capability you sell: 'How local businesses get found on ChatGPT,' "
    "'Why agencies hide their pricing,' 'What 24-hour delivery actually means.'",
])
h2("Test ideas (only after analytics work)")
bullets([
    "Hero A (outcome) vs C (risk-reversal). Metric: free-preview requests.",
    "Guarantee in hero vs. only at pricing table.",
    "Booking-first vs. form-first for high-intent traffic.",
    "Proof bar placement: under hero vs. beside plans.",
])

# =====================================================================
h1("Part 3, Still open")
bullets([
    "Frank: business type + any real outcome, and his approval of the testimonial.",
    "Jetta Grove: confirm the 4.34M/163-keyword work is SEO/discovery so it backs the core promise (and add a screenshot).",
    "Guarantee: confirm you can reliably hit 24 hours before publishing that half.",
    "Reviews: which platform you'll collect on (Google recommended for local SEO double-duty).",
])
note("Everything here is honest and defensible as written. The only true 'proof' assets right now are Frank (real) and "
     "your Jetta Grove track record (yours), so those carry the credibility load until reviews and closed prototypes "
     "come in. That's exactly the bottleneck the teardown identified; this plan buys credibility the fastest honest way.")

out = str(pathlib.Path(__file__).parent / "Well-and-Good-Websites-Action-Plan-and-Copy-2026-06-28.docx")
doc.save(out)
print("Saved:", out)
