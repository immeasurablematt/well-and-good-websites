#!/usr/bin/env python3
"""Build the live-site copy audit (captured copy + dual-framework analysis) as a .docx."""
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
ACCENT = RGBColor(0x1A, 0x53, 0x4F)
MUTED = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xB0, 0x00, 0x00)
GREEN = RGBColor(0x1E, 0x7A, 0x33)


def title(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(24)

def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)

def body(t):
    return doc.add_paragraph(t)

def sub(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = True; r.font.color.rgb = MUTED; return p

def label(lbl, value):
    p = doc.add_paragraph(); r = p.add_run(f"{lbl}: "); r.bold = True; r.font.color.rgb = ACCENT
    p.add_run(value); return p

def quote(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
    r = p.add_run(t); r.italic = True; return p

def bullets(items):
    for it in items: doc.add_paragraph(it, style="List Bullet")

def verdict(word, text):
    p = doc.add_paragraph(); r = p.add_run(f"{word} "); r.bold = True
    r.font.color.rgb = GREEN if word.startswith(("STRONG", "GOOD", "KEEP")) else (RED if word.startswith(("GAP", "WEAK", "MISSING", "FIX")) else ACCENT)
    p.add_run(text); return p

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""; run = c.paragraphs[0].add_run(hd); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row): cells[i].text = val
    doc.add_paragraph("")
    return t

# =====================================================================
title("Well and Good Websites — Live Copy Audit")
sub("Live copy pulled from wellandgoodwebsites.ca on 2026-06-27 (rendered HTML, JS on). "
    "Analyzed against two frameworks: (A) Positioning & Angles, (B) Copywriting Fundamentals.")

h2("Executive summary")
bullets([
    "The live site is strong on three angles: Speed/Ease (24 hours), Enemy/Contrarian (anti-agency transparency), and Risk Reversal (free preview, cancel anytime). These are well-chosen for a jaded, crowded local market.",
    "Biggest gap: almost no social proof for Well and Good itself — no testimonials, no count of sites built, no named clients or owner story. The only proof on the page (\"4.9-star reputation\") belongs to a client, not to WGW. Framework B's testimonial principle and Framework A's Social Proof angle are essentially unused.",
    "Second gap: the hero leads with a product attribute (\"Custom sites in 24 hours\") rather than the customer's transformation (get found → get chosen → more calls). Framework B says the headline does 80% of the work and should pair outcome + timeframe/contrast.",
    "Third gap: no named mechanism. In a Stage-4 market, buyers are skeptical of every \"more customers\" promise; naming the HOW (a branded system) raises believability.",
    "The free-preview offer (\"I'll build you a free preview of your website, no commitment\") is the single most persuasive line on the site — and it is buried in the final CTA. Move it above the fold.",
])

# =====================================================================
h1("Part 1 — Live copy, captured verbatim")
sub("Exactly as rendered on 2026-06-27. Layout labels (Hero, Card, Plan) added for navigation; wording unchanged.")

# ---- Homepage ----
h2("Homepage — wellandgoodwebsites.ca/")
label("SEO title", "Custom Websites in 24 Hours + Local Growth Marketing | Well and Good Websites  (77 chars — over the ~60–70 limit)")
label("Meta description", "Custom websites live in 24 hours, plus the local SEO, Google Business Profile, and social that get you found on Google. From $99/month — no lock-in, published pricing.")
label("OG title", "Custom Websites in 24 Hours + Local Growth Marketing")
h3("Hero")
quote("Custom sites in 24 hours. At 80% less than agency pricing.")
quote("(H2) Local agencies don't publish their pricing. We have nothing to hide.")
h3("Comparison table")
table(["What you need", "Typical Niagara agency", "Well and Good"], [
    ["Website build", "$2,500–$8,000, weeks to launch", "From $497, live in 24 hours"],
    ["Local SEO", "$800–$2,500/month", "Included in your plan"],
    ["Social media", "$950–$3,000/month", "Included in your plan"],
    ["Everything, bundled", "$1,800–$5,500/month", "From $99/month"],
    ["Their pricing", "\"Request a quote\"", "Published, right here"],
    ["Commitment", "6–12 month contracts", "Cancel anytime"],
])
body("Get found on Google, ChatGPT, and social media by customers searching for what you do. Show up first, show up often, stay ahead of the competition.")
h3("Recent Work")
body("Every site is built for local search, so customers on Google and ChatGPT find you first.")
body("Service provider · Toronto — \"SEO-optimized so local customers find it first.\" Services, videos, photos, testimonials, and contact — organized so visitors can decide fast.")
body("Restaurant · Welland — \"Built around the real reasons people visit.\" Menu highlights, reviews, location details, and call-to-order actions organized for mobile.")
body("Auto repair · Welland — \"Built to turn a 4.9-star reputation into phone calls.\" Services, real Google reviews, and a call-first layout — organized so drivers reach the shop before second-guessing.")
body("Agency results, a fraction of the price.")
h3("How it works — Three steps. Live in 24 hours.")
body("Understand — I review your Google listing, reviews, services, photos, social presence, and local competitors.")
body("Build — Your website goes live in 24 hours, written and optimized for local search.")
body("Grow — Ongoing SEO, Google Business Profile, and social keep you showing up and bringing in customers.")
h3("Plans — Get more attention. Get more customers.")
body("Launch — Get found locally. + $497 setup · or $990/yr")
bullets(["One-page site, mobile, tap-to-call", "Google Business Profile management",
         "Local SEO foundation + citations", "Hosting, SSL, backups, monitoring", "4 social posts/month, 1 platform"])
body("Grow (Most popular) — More calls, bookings, visits, quotes. + $997 onboarding · or $3,990/yr")
bullets(["Everything in Launch, plus:", "Standard site with booking/quote flow",
         "Active local SEO + 2 content pages/mo", "Managed social: 2 platforms, ~12/mo", "Review management + monthly report"])
body("Dominate — Own your local market. + $1,497 onboarding · or $8,990/yr")
bullets(["Everything in Grow, plus:", "Premium site + positioning & messaging",
         "Aggressive SEO + 4 content pieces/mo", "Social: 3–4 platforms + weekly video", "Monthly strategy call + dashboard"])
body("(H2) The whole package, for less than agencies charge for SEO alone. Ad budget billed separately. Prices in CAD.")
h3("Final CTA")
quote("Ready to get found and win more customers? Tell me about your business and I'll build you a free preview of your website, no commitment. Prefer email? matt@wellandgoodwebsites.ca")

# ---- Niagara ----
h2("Web Design Niagara — /web-design-niagara/")
label("SEO title", "Web Design Niagara + Local SEO and Social Growth")
label("Meta description", "Web design and growth marketing for Niagara businesses. Get found on Google and ChatGPT with custom websites plus local SEO, Google Business Profile, and social. From $99/month, published pricing.")
h3("Hero (H1)")
quote("Get your Niagara business found and chosen.")
h3("Built around how Niagara customers choose.")
body("A good local website makes the next step obvious on a phone: services, reviews, hours, directions, and one tap to call or book.")
body("Local proof — Real services, photos, reviews, and location details so the website feels specific to your business, not a generic template.")
body("Mobile actions — Tap-to-call, booking, quote requests, hours, and directions where customers find them quickly.")
body("Found first — Optimized for local search so Niagara customers find you on Google and ChatGPT before your competition.")
h3("Website + growth plans.")
body("Your Niagara website plus SEO, Google Business Profile, social, and content, for less than agencies charge for SEO alone. Cancel anytime.")
body("Launch, $99/mo (plus $497 setup) · Grow, $399/mo (plus $997 setup, Most popular) · Dominate, $899/mo (plus $1,497 setup).")
h3("Prefer just a website?")
body("One-time builds, no recurring fee: Express $297, Starter $497, Standard $997, Premium $1,497. Add Care for $39/month, upgrade anytime.")
h3("Final CTA")
quote("Want more Niagara customers finding and choosing you? Send me your business name, current link if you have one, and the main action you want from customers. I'll build you a free preview of your website, no commitment.")

# ---- Affordable ----
h2("Affordable Website Design — /affordable-website-design/")
label("SEO title", "Affordable Website Design + Growth Marketing | Well and Good Websites")
label("Meta description", "Affordable custom websites for Niagara and GTA businesses, plus the SEO, Google Business Profile, and social media growth that gets you found on Google and ChatGPT. From $99/month, published pricing.")
h3("Hero (H1)")
quote("A custom website, and the growth that fills it.")
h3("Affordable should still mean custom.")
body("Not a cheap-looking website. A clean one built around your real business and optimized to be found, without agency overhead.")
body("Real business details — Services, photos, hours, location, reviews, and the main customer action come before decoration.")
body("Clear pricing — Plans are visible from the start, so you can decide what level of help makes sense.")
body("Built to be found — Every website ships optimized for local search on Google and ChatGPT.")
h3("Plans + 'Prefer just a website?'")
body("Same Launch/Grow/Dominate plans. One-time builds: Express $297, Starter $497, Standard $997, Premium $1,497. Add Care for $39/month.")
h3("Final CTA")
quote("Want an affordable website that actually brings in customers? Send me your business name and current link, if you have one. I'll build you a free preview of your website, no commitment.")

# ---- One-page ----
h2("One-Page Websites — /one-page-websites/")
label("SEO title", "One-Page Websites for Local Businesses | Well and Good Websites")
label("Meta description", "Custom one-page websites for local businesses, built to be found on Google and ChatGPT and to turn visitors into calls and bookings. Available bundled with growth from $99/month.")
h3("Hero (H1)")
quote("One-page websites that make the next step obvious.")
h3("When a one-page website works well.")
body("Best when you have one clear offer, one service area, and a few key customer actions. Restaurants and shops; Trades and services; Appointment businesses.")
h3("What the page needs to include.")
body("Trust (reviews, real photos, recent work, owner note) · Search clarity (services, page title, headings, structured details) · Action (booking/quote, directions/hours, contact/social).")
h3("Final CTA")
quote("Need a simple one-page website for your business? Send me your business name, what customers need to do next, and any links. I'll build you a free preview of your website, no commitment.")

# =====================================================================
h1("Part 2 — Framework A: Positioning & Angles")
sub("The skill's 4-step process, Schwartz sophistication, 8 angle generators, the 5-point test, and a set of angle options.")

h2("Step 1 — Transformation: what changes after?")
verdict("PARTIAL.", "The site names the destination (\"get found,\" \"win more customers\") but stops one level short of the emotional/financial payoff. The before→after is implied, not dramatized. There is no \"from invisible to chosen\" arc and no owner-felt outcome (fewer slow weeks, a full calendar).")
body("Recommendation: state the after-state explicitly near the hero — e.g., \"Go from invisible on Google to the first name a local customer calls.\"")

h2("Step 2 — Competitive: map the alternatives")
verdict("STRONG.", "The comparison table is the best asset on the site. It maps the real alternative (local agencies) across six axes with specific dollar ranges, and wins on every one: price, inclusion, transparency, and commitment. This is textbook competitive positioning.")
body("Watch-out: the table only maps agencies. The other two alternatives a local owner considers — DIY builders (Wix/Squarespace) and doing nothing — are not addressed. Add a line that positions against DIY (\"done for you, not another tool to learn\").")

h2("Step 3 — Mechanism: name the HOW")
verdict("GAP.", "There is no named mechanism. \"Live in 24 hours\" and \"published pricing\" are claims, not an explained method. In a skeptical market, an unexplained 24-hour build invites the question \"how — and is it any good?\" The how (modern/AI tooling) is deliberately not shown; that's fine, but the site should still brand the system.")
body("Recommendation: name and trademark-style the process — e.g., \"The 24-Hour Local Launch\" or \"The Found-First Build\" — so the speed reads as a repeatable system, not a corner-cut.")

h2("Step 4 — Market sophistication (Schwartz)")
verdict("STAGE 4 (jaded).", "Local web design / SEO is a crowded, claim-fatigued market — every agency promises \"more customers\" and \"#1 on Google.\" Stage 4 rewards identity and a fresh mechanism over louder claims; Stage 5 rewards exclusivity/identity.")
body("The site is already playing a smart Stage-4 hand: it leads with an enemy (secretive agencies) and an identity (the transparent, no-lock-in alternative) rather than a bare \"best websites\" claim. To push further, lean harder into identity (\"the anti-agency for owner-operators\") and add the named mechanism from Step 3.")

h2("The 8 angle generators — which are working")
table(["Angle", "Used on site?", "Evidence / opportunity"], [
    ["Contrarian (challenge beliefs)", "Yes — strong", "\"Local agencies don't publish their pricing. We have nothing to hide.\""],
    ["Unique Mechanism (lead with HOW)", "No", "No named method; 24h is a claim, not a mechanism. Biggest untapped angle."],
    ["Transformation (before vs after)", "Partial", "\"Get found... win more customers\" — destination named, arc not dramatized."],
    ["Enemy (rally vs villain)", "Yes — strong", "The villain is the opaque, lock-in agency. Carried through the table + hero."],
    ["Speed / Ease (compress time)", "Yes — strong", "\"Custom sites in 24 hours,\" \"Three steps. Live in 24 hours.\""],
    ["Specificity (hyper-targeted avatar)", "Partial", "Niagara/Welland + verticals (auto, restaurant) appear, but hero speaks to \"businesses\" broadly."],
    ["Social Proof (lead with evidence)", "MISSING", "No WGW testimonials, no site count, no named clients. Only a client's \"4.9-star\" reputation."],
    ["Risk Reversal (guarantee as headline)", "Yes, but buried", "\"Free preview, no commitment\" + \"cancel anytime\" — excellent, but sits in the footer CTA."],
])

h2("The 5-point angle test — applied to the current hero")
sub("Hero tested: \"Custom sites in 24 hours. At 80% less than agency pricing.\"")
table(["Test", "Pass?", "Notes"], [
    ["1. Specific?", "Yes", "24 hours and 80% are concrete and memorable."],
    ["2. Differentiated?", "Partial", "Speed + price are differentiated locally, but \"custom sites\" is generic; many claim it."],
    ["3. Believable?", "At risk", "\"24 hours\" + \"80% less\" can read as too-good-to-be-true with no mechanism or proof to anchor it."],
    ["4. Relevant to THIS audience?", "Partial", "Speaks to price-shoppers; less to owners whose real pain is \"no one finds me.\""],
    ["5. Does it lead somewhere?", "Yes", "Flows into the comparison table and plans."],
])
verdict("READ:", "The hero passes on specificity but is weakest on believability (no mechanism/proof) and on speaking to the deeper want (being chosen, not just being cheap and fast). Both are fixable with the angle options below.")

h2("Angle options (the skill's required output: 3–5 options)")
sub("Each: the angle, the psychology it pulls, a headline direction, and when to lead with it.")

h3("Option 1 — Enemy / Transparency (the anti-agency)")
label("Psychology", "Distrust of agencies; relief at being treated like an adult. Rallying against a shared villain builds fast allegiance.")
label("Headline direction", "\"Everything local agencies hide — the price, the timeline, the contract — out in the open.\"")
label("When to use", "Lead angle for cold local owners who've been quoted by (or burned by) an agency. Strongest fit for the current audience.")

h3("Option 2 — Transformation (invisible → chosen)")
label("Psychology", "Status and survival: being found and picked over the competitor down the street.")
label("Headline direction", "\"Go from invisible on Google to the first call a local customer makes.\"")
label("When to use", "When the visitor's pain is lack of customers, not price. Good for SEO/organic traffic searching problems, not prices.")

h3("Option 3 — Risk Reversal (see it before you pay)")
label("Psychology", "Loss aversion removed: zero risk to look. The free preview makes the first step free and concrete.")
label("Headline direction", "\"See your new website before you pay a cent.\"")
label("When to use", "Paid traffic and skeptical price-sensitive visitors. Highest-converting first step; deserves above-the-fold placement.")

h3("Option 4 — Unique Mechanism (name the system)")
label("Psychology", "Believability: a named method makes 24 hours sound engineered, not rushed.")
label("Headline direction", "\"The 24-Hour Local Launch: your site live tomorrow, built to be found.\"")
label("When to use", "Pair with any hero to shore up the believability gap from the 5-point test.")

h3("Option 5 — Bundle / One throat to choke")
label("Psychology", "Simplicity and relief from vendor-juggling; one bill, one person.")
label("Headline direction", "\"Website, SEO, and social — one person, one bill, for less than agencies charge for SEO alone.\"")
label("When to use", "For owners already paying multiple vendors or overwhelmed by DIY. This is the plans' real differentiator.")

# =====================================================================
h1("Part 3 — Framework B: Copywriting Fundamentals")
sub("Audited principle by principle, with specific lines from the live site.")

h2("Headlines do 80% of the work")
body("Formula from the framework: action verb + specific outcome + timeframe or contrast.")
verdict("MIXED.", "The homepage hero (\"Custom sites in 24 hours. At 80% less than agency pricing.\") has timeframe and contrast but no action verb and no customer outcome — it describes the product, not the result. The Niagara hero (\"Get your Niagara business found and chosen.\") is the best on the site: action verb + outcome + audience. The one-page hero (\"One-page websites that make the next step obvious.\") is descriptive and the weakest.")
body("Write five hero variants and test. Five starters for the homepage:")
bullets([
    "Get found on Google and chosen by local customers — with a site live in 24 hours.",
    "Your business online tomorrow, found first, for 80% less than an agency.",
    "See your new website before you pay a cent — live in 24 hours.",
    "The whole growth system — site, SEO, social — for less than agencies charge for SEO alone.",
    "Stop losing local searches to the competitor who showed up first.",
])

h2("Opening lines: earn the second sentence")
verdict("STRONG.", "\"Local agencies don't publish their pricing. We have nothing to hide.\" is a direct challenge + confession — exactly the framework's prescription, and it avoids the \"In today's fast-paced world\" trap entirely. Keep it.")

h2("Pain quantification: do the math")
verdict("HALF DONE.", "The comparison table quantifies the competitor's cost ($1,800–$5,500/mo) brilliantly. But the customer's cost of inaction is never quantified. Add the other side of the math:")
quote("Every month you're not showing up, the three competitors who are split the calls you didn't get.")

h2("The 'so what' chain: go three levels deep")
verdict("STOPS TOO HIGH.", "\"Live in 24 hours\" stops at the feature. Push it down to the emotional/financial floor:")
bullets([
    "Live in 24 hours →",
    "you're on Google this week, not in six weeks →",
    "calls start while your competitor is still waiting on their agency →",
    "you stop watching jobs go to the shop down the street.",
])
body("Write the hero from the bottom of that chain, not the top.")

h2("Open loops: keep them reading")
verdict("FEW.", "The page is efficient but rarely teases forward. One or two loops would pull readers down — e.g., after the table: \"There's one reason I can charge this little. More on that below.\" (then reveal the bundle/efficiency, not the cost structure).")

h2("Rhythm: short and long, not uniform")
verdict("GOOD.", "Sentence length varies and the copy reads human, not AI-uniform. No change needed. Minor: the recent-work cards lean on em-dashes (\"contact — organized,\" \"call-first layout — organized\"); the framework flags excessive em-dashes, so vary the punctuation in a couple of spots.")

h2("Testimonials: the missing engine")
verdict("MISSING — top priority.", "The framework's testimonial formula (before state + action + specific outcome + timeframe + emotion) cannot be applied because there are zero testimonials for Well and Good. The only proof on the site (\"a 4.9-star reputation\") is the client's, not yours. This is the single biggest believability gap and it undercuts the bold 24h/80% claims.")
body("Action: collect 2–3 real client lines in the framework's shape, e.g. \"I'd put off a website for two years. Matt had mine live the next day and I booked three jobs that week.\" Until real ones exist, do not fabricate — instead substitute concrete proof you can stand behind: number of sites launched, a named local client with permission, or a screenshot of a real result.")

h2("Avoid-list compliance")
verdict("CLEAN.", "No \"utilize,\" \"delve,\" \"leverage,\" \"game-changer,\" or \"Let's dive in.\" \"Dominate\" / \"Own your local market\" are aggressive but acceptable as plan names. Em-dash density is the only minor flag (see Rhythm).")

# =====================================================================
h1("Part 4 — Prioritized recommendations")

h2("P1 — Highest impact (do first)")
bullets([
    "Add real social proof. 2–3 testimonials in the before→after→outcome→emotion shape, plus a concrete number (\"X local sites launched\"). Place one near the hero and one beside the plans. This is the #1 fix — it makes the 24h/80% claims believable.",
    "Move the free preview above the fold. \"See your new website before you pay a cent — no commitment\" is the strongest line on the site and it's buried in the footer. Make it a hero sub-CTA.",
    "Rewrite the hero from the bottom of the so-what chain. Test an outcome-led variant (\"Get found and chosen — site live in 24 hours\") against the current attribute-led one. Write 5, test the top 2.",
])
h2("P2 — High impact")
bullets([
    "Name the mechanism (\"The 24-Hour Local Launch\" / \"Found-First Build\") to close the believability gap on speed.",
    "Quantify the cost of inaction near the comparison table (the customer's side of the math).",
    "Add a DIY-builder row or line to the comparison (position against Wix/Squarespace and 'doing nothing', not just agencies).",
    "Push the identity angle harder: 'the anti-agency for owner-operators who'd rather run their business than their website.'",
])
h2("P3 — Polish & technical")
bullets([
    "Shorten the homepage SEO title (77 chars → under ~65) so it doesn't truncate in search results.",
    "Add image alt text (the audit flags no_image_alt) — accessibility + local image SEO.",
    "Homepage content is thin (~527 words, flagged low content rate); the testimonials and mechanism copy above will help.",
    "Reduce em-dash density in the recent-work cards.",
    "Add 1–2 forward-teasing open loops to pull readers down the page.",
])

h2("Test ideas (A/B, don't just swap)")
bullets([
    "Hero: attribute-led (current) vs. outcome-led vs. free-preview-led. Metric: preview requests.",
    "Free preview as primary CTA vs. 'Get my plan'. Metric: lead submissions.",
    "Mechanism-named hero vs. unnamed. Metric: scroll depth + conversion (believability proxy).",
    "Social-proof block placement: near hero vs. near plans. Metric: plan-section conversion.",
])

h2("Context flags / where I need input")
bullets([
    "Testimonials: do you have real client results and permission to quote them? Needed before any proof copy ships — nothing will be fabricated.",
    "The 24-hour + 80%-less claims should be defensible if challenged; confirm they hold for a typical build.",
    "Earlier open items still stand: the footer name/link (Frank vs. Matthew) and an owner story for the About/identity angle.",
])

out = str(pathlib.Path(__file__).parent / "Well-and-Good-Websites-Live-Copy-Audit-2026-06-27.docx")
doc.save(out)
print("Saved:", out)
