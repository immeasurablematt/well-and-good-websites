#!/usr/bin/env python3
"""Build the Well and Good Websites revised copy deck as a formatted .docx."""
import pathlib

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

ACCENT = RGBColor(0x1A, 0x53, 0x4F)  # deep teal for labels


def title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def body(text):
    return doc.add_paragraph(text)


def label(lbl, value):
    """Bold label: value on one line (for SEO fields)."""
    p = doc.add_paragraph()
    r = p.add_run(f"{lbl}: ")
    r.bold = True
    r.font.color.rgb = ACCENT
    p.add_run(value)
    return p


def cta(primary, micro=None, secondary=None):
    p = doc.add_paragraph()
    r = p.add_run(f"[ {primary} ]")
    r.bold = True
    r.font.size = Pt(12)
    # bold accent-colored text stands in for a button
    r.font.color.rgb = ACCENT
    if secondary:
        sp = doc.add_paragraph()
        sr = sp.add_run(f"Secondary: {secondary}")
        sr.italic = True
    if micro:
        mp = doc.add_paragraph()
        mr = mp.add_run(micro)
        mr.italic = True
        mr.font.size = Pt(9)


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def flag(text):
    p = doc.add_paragraph()
    r = p.add_run(f"⚑ {text}")
    r.bold = True
    r.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)


def card(name, who, feats, popular=False):
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(13)
    if popular:
        b = p.add_run("   ★ MOST POPULAR")
        b.bold = True
        b.font.color.rgb = ACCENT
    if who:
        wp = doc.add_paragraph()
        wr = wp.add_run(who)
        wr.italic = True
    bullets(feats)


# =====================================================================
title("Well and Good Websites: Website Copy Deck")
sub = doc.add_paragraph()
sr = sub.add_run("Revised 2026-06-16 · updated positioning, transparent pricing, and optional care plans (recurring)")
sr.italic = True
sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

h2("What changed in this revision")
bullets([
    "Primary promise dropped “local”: “A custom website for your business, live in 48 hours.”",
    "Audience broadened from “local businesses” to small and medium businesses, serving both leads and existing customers.",
    "Positioning leads with “custom websites”; one-page is presented as a format (still its own page + SEO term), not the headline identity.",
    "Niagara/Toronto local SEO targeting kept intact (organic focus unchanged).",
    "NEW: optional recurring care plans (Essential / Growth / Pro) added to the homepage and woven into the FAQ, pricing left as $XX placeholders until you decide.",
    "NEW: transparency wedge (“every price is on the page”) and “no lock-in / cancel anytime” folded through all pages, angles drawn from the competitor pricing analysis.",
    "Kept OFF the site by design: margins, cost structure, “AI-built” as a lead message, and hard competitor price multiples (accuracy/legal risk).",
    "Carried over from the prior pass: primary CTA “Get my website quote”, prices on package cards, “Most popular” on Standard, risk-reversal FAQ, trust line.",
])
flag("Two items still need your input, see the red flags inline: the footer name/link (Frank vs. Matthew) and a real proof number for the trust line. Nothing fabricated.")

h2("Positioning (reference)")
label("Primary promise", "A custom website for your business, live in 48 hours.")
label("Positioning", "Practical custom websites for small and medium businesses that need calls, bookings, and visits, from leads and customers who need quotes, service details, hours, and directions.")
label("Organic focus", "web design Niagara, affordable website design, affordable website design for small business, one-page websites, local website design for Niagara and Toronto businesses.")
label("Pages covered", "Homepage, Web Design Niagara, Affordable Website Design, One-Page Websites, robots.txt, sitemap.xml.")

# =====================================================================
h1("Homepage")
label("Live path", "https://wellandgoodwebsites.ca/")
label("SEO title", "Affordable Custom Websites for Small Businesses | Well and Good Websites")
label("Meta description", "Affordable custom websites for small and medium businesses in Niagara and Toronto. Well and Good Websites builds fast, practical sites that help leads and customers call, book, visit, or request a quote.")
label("Open Graph title", "Affordable Custom Websites for Small and Medium Businesses")
label("Open Graph description", "Custom websites for Niagara and Toronto businesses, live in 48 hours and built to turn visitors into calls, bookings, visits, and quote requests.")

h2("Hero")
body("Well and Good")
body("WEBSITES")
h3("A custom website for your business, live in 48 hours.")
body("Custom websites for small and medium businesses in Niagara and Toronto, built so leads and customers can call, book, visit, or get a quote, and find your hours and directions fast.")
cta("Get my website quote",
    micro="Just your business name and current link. I’ll reply with the simplest next step.",
    secondary="See recent builds")
tp = doc.add_paragraph()
tr = tp.add_run("Built by Matthew Baggetta for Niagara and Toronto businesses. ")
tr.bold = True
flag("[PROOF NUMBER NEEDED, e.g. “12+ sites launched” or a review count. Left blank rather than invented.]")

h2("Recent builds")
body("Focused sites that put the most useful business details up front, so visitors understand the offer and take the next step quickly.")
h3("A service business with a real home base.")
body("Services, video, photos, testimonials, and contact details arranged so leads decide fast.")
cta("View live site", secondary="See packages")
h3("A restaurant built around why people actually visit.")
body("Menu highlights, reviews, location, and call-to-order actions organized for phones.")
cta("View live site", secondary="See packages")

h2("Simple packages")
body("Built for owner-operated small and medium businesses, auto shops, barbers, salons, tattoo artists, bakeries, contractors, performers, and appointment-based services, that want a real website without agency prices, long timelines, or cookie-cutter templates.")
body("Every price is on the page. Most agencies make you request a quote first, here you can decide before you ever call.")
card("Starter, $497", "Just need to exist and be findable.", [
    "Custom website", "Mobile responsive design", "Tap-to-call and contact section",
    "Hours, map, reviews, services", "Basic search setup"])
card("Standard, $997", "Want the site to actively bring in calls and bookings.", [
    "Everything in Starter", "Photo gallery or portfolio section",
    "Booking or quote-request flow", "Review highlights", "Google Business Profile support"],
    popular=True)
card("Premium, $1,497", "Want the site plus a system that turns reviews into customers.", [
    "Everything in Standard", "Custom copywriting and positioning",
    "Before/after or menu/order section", "Review funnel page", "Priority launch support"])

h2("Live, and looked after.")
body("Your site is yours, with no lock-in. An optional care plan keeps it hosted, secure, backed up, and current, and small edits are included, not billed by the quarter-hour. Every price is published; cancel anytime, or pay yearly for two months free.")
care_note = "[CARE-PLAN PRICING NOT FINAL, $XX / $XXX are placeholders. Recommendation on the table: $19 / $39 / $79 monthly. Drop final numbers in here and in the FAQ once decided.]"
card("Essential, $XX/mo ($XXX/yr)", "Keep the site online, secure, and current.", [
    "Fast, secure hosting and SSL", "Domain and uptime monitoring", "Automatic backups",
    "One content edit a month (hours, photos, text)", "Email and Telegram support"])
card("Growth, $XX/mo ($XXX/yr)", "For businesses whose details change often.", [
    "Everything in Essential", "Unlimited small content edits, 24–48 hour turnaround",
    "Quarterly booking and link check", "One seasonal refresh a year"], popular=True)
card("Pro, $XX/mo ($XXX/yr)", "Your site plus an active local presence.", [
    "Everything in Growth", "Google Business Profile management", "Review-funnel upkeep",
    "Light monthly local SEO with a plain-English performance note", "Priority turnaround"])
body("New builds get the first month free on the matching plan. Prefer to handle your own updates? You can, single edits start at $XX, no plan required.")
flag(care_note)

h2("Website design for local searches")
body("Each site is built around the details customers actually look for: what you do, where you are, when you are open, what it costs, what customers say, and how to take the next step.")
h3("Niagara, web design for Niagara businesses")
body("Custom websites for restaurants, trades, shops, salons, barbers, contractors, and appointment-based businesses across Niagara.")
h3("Packages, affordable website design")
body("Clear $497, $997, and $1,497 website packages for small and medium businesses that need a useful site without a slow agency process.")
h3("Format, one-page websites that convert")
body("When it fits, a focused single-page site built around calls, bookings, quote requests, directions, reviews, photos, services, and menus.")

h2("Clear process")
body("No long discovery maze. I turn the useful pieces your business already has into a custom website that is easy to find, trust, and contact.")
h3("1. Understand")
body("I review your Google listing, current links, reviews, services, photos, and nearby competitors.")
h3("2. Shape")
body("I map the page around your business, your customers, and the actions that matter: call, visit, book, or request a quote.")
h3("3. Build")
body("I collect the right photos, details, and approvals, then build the site with custom copy, layout, and mobile-first design.")
h3("4. Launch")
body("The site goes live within 48 hours after the details are approved, giving your Google profile, referrals, and follow-ups one clear place to point.")

h2("Why this works")
body("Many small and medium businesses already have the hard part: real customers, strong reviews, and a reputation. A clear website turns that trust into a faster next step.")
h3("Custom beats generic.")
body("Your site should use your real services, reviews, location, photos, and customer questions, not a template that could belong to anyone.")
h3("Mobile visitors should not have to think.")
body("Directions, hours, and tap-to-call should be obvious on a phone, not buried.")
h3("Your reviews should do more selling.")
body("The best sites make your real customer words easier to see, trust, and act on.")

h2("Questions owners ask")
body("Short answers for businesses comparing a custom website with templates, free website builders, or a traditional agency project.")
h3("How much does a small business website cost?")
body("Every price is published up front, no “request a quote” maze. The build is a one-time $497, $997, or $1,497, and the right fit depends on how much copy, photography, booking flow, review support, and launch help you need. Keeping the site online afterward is an optional care plan, also priced openly.")
h3("Can a one-page website rank on Google?")
body("A one-page website can rank for focused local searches when it has clear services, location language, fast loading, useful headings, strong business details, and links from real local profiles.")
h3("What do you need to launch in 48 hours?")
body("The fastest launches start with your business name, services, location, hours, photos, reviews, preferred contact method, and any existing Google Business Profile or social links.")
h3("Do you help with Google Business Profile?")
body("Yes. The Standard and Premium packages include support for connecting the website to your Google Business Profile so customers have one clear place to call, visit, book, or request a quote.")
h3("Is this better than Wix or Squarespace?")
body("Website builders can work if you have the time and confidence to write, design, structure, and launch the site yourself. Well and Good is done for you, like the bigger online website services, but without a large setup fee or a $200-a-month plan, and built fast around your real business.")
h3("What does it cost to keep my site online?")
body("Hosting, security, backups, and small updates are handled through an optional care plan, published up front from $XX a month, no setup fee, no lock-in. Prefer not to subscribe? You can pay per edit instead.")
h3("Can I change my site after it’s live?")
body("Yes. Send the change by message, hours, a photo, a price, a new section, and I make it for you. On the Growth and Pro care plans, small edits are included instead of billed by the quarter-hour.")
h3("Is there a contract or lock-in?")
body("No. Care plans are month to month, you can cancel anytime, and you own your domain. Pay yearly if you’d like two months free, but you are never locked in.")
h3("What if I don’t like it?")
body("I share a first draft before launch and revise it until it’s right for your business. You approve the details before the site goes live.")
flag("Confirm this revision-policy wording matches what you can actually promise.")

h2("Final CTA")
h3("Ready to make your business easier to choose?")
body("Send me your business name and current web link, if you have one. I’ll reply with the simplest useful next step for getting a custom site live quickly.")
cta("Get my website quote")
body("Connect on LinkedIn")
flag("[CONFIRM NAME/LINK, doc currently shows “FrankBaggetta.ca” but the builder is Matthew Baggetta.]")
body("X")
body("Back to top")

# =====================================================================
h1("Web Design Niagara Page")
label("Live path", "https://wellandgoodwebsites.ca/web-design-niagara/")
label("SEO title", "Web Design Niagara | Affordable Local Business Websites")
label("Meta description", "Affordable web design for Niagara small and medium businesses. Well and Good Websites builds custom sites for restaurants, trades, shops, salons, contractors, and local services.")
label("Open Graph title", "Web Design Niagara")
label("Open Graph description", "Affordable custom websites for Niagara businesses that need more calls, bookings, visits, and quote requests.")

h2("Hero")
body("Well and Good Websites")
body("Web design Niagara")
h3("Custom websites for Niagara businesses, built to win more local calls and bookings.")
body("Well and Good Websites builds focused sites for businesses across Niagara that need leads and customers to call, visit, book, or request a quote.")
cta("Get my website quote", secondary="See recent builds")
body("Designed for restaurants, trades, shops, barbers, salons, contractors, performers, and appointment-based services.")

h2("Built around how Niagara customers choose.")
body("A useful local website should make the next step obvious on a phone. That means services, photos, reviews, hours, directions, booking links, quote requests, and calls all working together.")
h3("Local proof")
body("Use real services, photos, reviews, and location details so the site feels specific to your business, not a generic template.")
h3("Mobile actions")
body("Put tap-to-call, booking, quote requests, hours, and directions where customers can find them quickly.")
h3("Basic search setup")
body("Connect the site to clear page titles, descriptions, headings, schema, analytics, and your Google Business Profile.")

h2("What can be included.")
body("The first build is usually a focused one-page website. Extra pages can be added when there is a clear search reason.")
card("Starter, $497", "", ["Custom one-page website", "Services, hours, map, reviews", "Mobile responsive design"])
card("Standard, $997", "", ["Everything in Starter", "Gallery or portfolio section", "Booking or quote-request flow", "Google Business Profile support"], popular=True)
card("Premium, $1,497", "", ["Everything in Standard", "Custom copywriting and positioning", "Review funnel page", "Priority launch support"])

h2("Need a Niagara business website that is simple to launch?")
body("Send your business name, current link if you have one, and the main action you want from customers. I’ll reply with the simplest useful next step.")
body("Prices are published up front, and the optional care plan keeps your site current with no lock-in.")
cta("Get my website quote")

# =====================================================================
h1("Affordable Website Design Page")
label("Live path", "https://wellandgoodwebsites.ca/affordable-website-design/")
label("SEO title", "Affordable Website Design for Small Businesses | Well and Good")
label("Meta description", "Affordable website design for small and medium businesses in Niagara and Toronto. Clear $497, $997, and $1,497 packages for custom websites built fast.")
label("Open Graph title", "Affordable Website Design for Small Businesses")
label("Open Graph description", "Clear custom website packages for businesses that need calls, bookings, visits, and quote requests.")

h2("Hero")
body("Well and Good Websites")
body("Affordable website design")
h3("Custom websites without agency prices or slow timelines.")
body("Well and Good Websites builds affordable custom websites for small and medium businesses in Niagara and Toronto, with clear packages and practical launch support.")
cta("Get my website quote", secondary="Compare packages")
body("Starter websites begin at $497.")
body("Built around the details customers need before they call, visit, book, or request a quote.")

h2("Affordable should still mean custom.")
body("The goal is not a cheap-looking site. The goal is to use the strongest parts of your real business and launch a clean website without unnecessary agency overhead.")
h3("Real business details")
body("Services, photos, hours, location, reviews, and the main customer action come before decoration.")
h3("Clear pricing")
body("Every price is on the page, builds and care plans both. No “request a quote” maze, no surprise invoice. Most agencies nearby hide their pricing; here you decide before you ever call.")
h3("Fast launch path")
body("Once the required details are approved, the site is built around a 48-hour launch target.")

h2("Website packages.")
body("Choose the level that matches how much structure, copy, and launch support your business needs.")
card("Starter: $497", "", ["Custom website", "Mobile responsive layout", "Hours, map, reviews, services", "Basic search setup"])
card("Standard: $997", "", ["Everything in Starter", "Photo gallery or portfolio", "Booking or quote-request flow", "Google Business Profile support"], popular=True)
card("Premium: $1,497", "", ["Everything in Standard", "Custom copywriting", "Review funnel page", "Priority launch support"])

body("Keeping the site online afterward is an optional care plan from $XX a month, published up front, cancel anytime, and you own your domain. No lock-in, ever.")
flag("[CARE-PLAN PRICING placeholder $XX, keep in sync with the homepage care section.]")

h2("Want an affordable website that still feels like your business?")
body("Send your business name and current web link, if you have one. I’ll reply with the simplest useful next step.")
cta("Get my website quote")

# =====================================================================
h1("One-Page Websites Page")
label("Live path", "https://wellandgoodwebsites.ca/one-page-websites/")
label("SEO title", "One-Page Websites for Local Businesses | Well and Good")
label("Meta description", "Custom one-page websites for businesses that need a focused site for calls, bookings, visits, quote requests, reviews, services, hours, and directions.")
label("Open Graph title", "One-Page Websites for Local Businesses")
label("Open Graph description", "Focused single-page websites built around calls, bookings, quote requests, reviews, services, hours, and directions.")

h2("Hero")
body("Well and Good Websites")
body("One-page website design")
h3("One-page websites that make the next step obvious.")
body("A focused single-page website can be the right fit when customers need clear services, photos, reviews, hours, directions, and one simple way to act.")
cta("Get my website quote", secondary="See recent builds")
body("Built for mobile visitors who want to know what you do, whether they trust you, and how to contact you.")

h2("When a one-page website works well.")
body("Single-page websites are best when the business has one clear offer, one local service area, and a few important customer actions.")
h3("Restaurants and shops")
body("Menu highlights, reviews, photos, location, hours, and call-to-order or directions can live on one focused page.")
h3("Trades and services")
body("Services, proof, quote requests, phone number, service area, and recent work can be arranged for quick decisions.")
h3("Appointment businesses")
body("Booking links, services, prices, reviews, photos, and policies can be grouped into one mobile-friendly flow.")

h2("What the page needs to include.")
body("The best one-page websites are not short because details are missing. They are short because the useful details are organized tightly.")
h3("Trust")
bullets(["Reviews and testimonials", "Real photos", "Recent work or menu items", "Business story or owner note"])
h3("Search clarity")
bullets(["Services and service area", "Helpful page title", "Useful headings", "Structured business details"])
h3("Action")
bullets(["Tap-to-call", "Booking or quote request", "Directions and hours", "Contact and social links"])

h2("Need a simple one-page website for your business?")
body("Send your business name, what customers need to do next, and any current links. I’ll map the simplest useful website for you.")
body("Clear prices up front, and an optional care plan to keep it current, month to month, no lock-in.")
cta("Get my website quote")

# =====================================================================
h1("Technical SEO Assets")
h2("robots.txt")
for line in ["User-agent: *", "Allow: /", "Sitemap: https://wellandgoodwebsites.ca/sitemap.xml"]:
    p = doc.add_paragraph(line)
    p.style = doc.styles["Normal"]
    p.runs[0].font.name = "Consolas"
h2("sitemap.xml URLs")
bullets([
    "https://wellandgoodwebsites.ca/",
    "https://wellandgoodwebsites.ca/web-design-niagara/",
    "https://wellandgoodwebsites.ca/affordable-website-design/",
    "https://wellandgoodwebsites.ca/one-page-websites/",
])

out = str(pathlib.Path(__file__).parent / "Well-and-Good-Websites-Copy-Deck-Revised-2026-06-16.docx")
doc.save(out)
print("Saved:", out)
