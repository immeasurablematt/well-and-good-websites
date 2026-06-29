#!/usr/bin/env python3
"""Build the Actions Plan (non-copy items) as a .docx. Copy edits live in the tracker."""
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11)
ACCENT = RGBColor(0x1A, 0x53, 0x4F); MUTED = RGBColor(0x55, 0x55, 0x55)

def title(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(24)
def sub(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = True; r.font.color.rgb = MUTED
def h2(t): doc.add_heading(t, level=2)
def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""; run = c.paragraphs[0].add_run(hd); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = v
    doc.add_paragraph("")

title("Well and Good Websites: Actions Plan")
sub("Non-copy items: structural, page-build, and technical work. Copy edits are tracked separately in "
    "Well-and-Good-Websites-Copy-Recommendations-Tracker. Pulled from the live-site passes, 2026-06-29.")

h2("Build / structural")
table(["Action", "Where", "Why", "Priority", "Status"], [
    ["Repeat the primary CTA at each decision point (after the table, testimonial, About, pricing).",
     "Homepage", "Long page loses visitors between CTAs.", "High", "OPEN"],
    ["Add an instant \"Book a 15-min call\" option (Cal.com) beside the form.",
     "Homepage", "High-intent visitors can act now instead of waiting for a reply.", "High", "OPEN"],
    ["Add a short FAQ section (what you need from me / is a 24-hour site any good / can I cancel / do I own it).",
     "Homepage", "Concentrates objection handling right before the CTA.", "Med", "OPEN"],
    ["Add Frank's testimonial plus a \"Built by Matt, here in Welland\" proof block.",
     "Niagara + Affordable pages", "These SEO pages convert worse than the homepage; they lack any proof.", "High", "OPEN"],
    ["Move the guarantee line into the hero (not just the bottom CTA).",
     "All sub-pages", "Risk reversal belongs above the fold.", "Med", "OPEN"],
    ["Relabel or move the github.io prototypes (\"Concept build\" or real subdomains); open live links in a new tab.",
     "Homepage Recent Work", "Spec builds read as fake client work and leak visitors off-site.", "High", "VERIFY"],
])

h2("Trust signals")
table(["Action", "Why", "Priority", "Status"], [
    ["Collect 5 to 10 Google reviews and embed the star badge near the hero (ask Frank first).",
     "Competitors lead with 5.0-star ratings; you have none yet.", "Med", "OPEN"],
])

h2("Technical")
table(["Action", "Why", "Priority", "Status"], [
    ["Add descriptive alt text to all images.", "Accessibility plus local image SEO (flagged: no_image_alt).", "Low", "OPEN"],
    ["Verify GA4 and Hotjar fire before running any A/B test.", "You cannot measure a test if events do not fire.", "High", "OPEN"],
])

out = str(pathlib.Path(__file__).parent / "Well-and-Good-Websites-Actions-Plan.docx")
doc.save(out)
print("Saved:", out)
