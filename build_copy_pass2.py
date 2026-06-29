#!/usr/bin/env python3
"""Build the second copy pass (copywriting then CRO) as a .docx."""
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11)
ACCENT = RGBColor(0x1A, 0x53, 0x4F); MUTED = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xB0, 0x00, 0x00); GREEN = RGBColor(0x1E, 0x7A, 0x33)

def title(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(24)
def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)
def body(t): return doc.add_paragraph(t)
def sub(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = True; r.font.color.rgb = MUTED; return p
def bullets(items):
    for it in items: doc.add_paragraph(it, style="List Bullet")
def cur(t):
    p = doc.add_paragraph(); r = p.add_run("Current:  "); r.bold = True; r.font.color.rgb = MUTED
    rr = p.add_run(t); rr.italic = True; return p
def rew(t):
    p = doc.add_paragraph(); r = p.add_run("Rewrite:  "); r.bold = True; r.font.color.rgb = ACCENT
    p.add_run(t); return p
def why(t):
    p = doc.add_paragraph(); r = p.add_run("Why:  "); r.bold = True; r.font.color.rgb = GREEN
    p.add_run(t); return p
def flag(t):
    p = doc.add_paragraph(); r = p.add_run(f"⚑ {t}"); r.bold = True; r.font.color.rgb = RED

# =====================================================================
title("Well and Good Websites — Copy Pass #2")
sub("Fresh pull of the live site, 2026-06-29. Copywriting pass first, then CRO pass — as requested. "
    "The site has clearly absorbed the last round: outcome-led H1, the 24-hour guarantee, Frank's testimonial with a "
    "real result, the About-Matt block, and the Jetta Grove proof are all live. So this is a refinement pass, not a rebuild.")

h2("Already shipped since last pass (good — leave these)")
bullets([
    "H1 is now outcome-led: \"Get found by local customers, and chosen before your competitors.\"",
    "Guarantee is live and repeated: \"Your site live in 24 hours, or your first month is free. Then I'll revise it until you're happy.\"",
    "Frank's testimonial is in, with a real, attributed result: \"Fifteen years online with zero leads. Three enquiries in his first month.\" (Search Console + GA4 noted — honest.)",
    "About-Matt block is live, first-person, with the Prometheus line and the Jetta Grove proof (4.34M impressions, 163 keywords, Page 1) linked.",
    "Eyebrow rewritten to the bundle idea; comparison table retained.",
])

# =====================================================================
h1("Part A — Copywriting pass")
sub("Skill: copywriting. Section-by-section, homepage. Each: current copy, a rewrite, and the principle behind it.")

h2("1. Hero headline")
cur("Get found by local customers, and chosen before your competitors.")
rew("Get found by local customers — and chosen before the competitor down the street. Live in 24 hours.")
why("Specificity over vagueness. The current H1 is benefit-led but could belong to any local marketing agency; it "
    "dropped the two things that made the old hero unmistakably you — the 24-hour speed and the 80%-less price. Pull one "
    "concrete element back into the headline (timeframe), and make the rival concrete (\"down the street\").")
body("Three more to test:")
bullets([
    "A — \"Get found on Google and chosen by local customers — your site live in 24 hours.\"",
    "B — \"The website an agency charges $5,000 for. Found-first, live in 24 hours, from $99/month.\"",
    "C — \"See your new website before you pay a cent — live in 24 hours.\" (leads with the free-preview hook)",
])

h2("2. Primary CTA button copy")
flag("The button label isn't in the page text I can read — confirm it. If it's generic (\"Get started\" / \"Request a website\"), change it.")
rew("Get my free preview")
why("CTA formula = action verb + what they get. Your strongest, lowest-risk asset is \"a free preview, no commitment\" — "
    "put that value in the button itself, not just in the paragraph above it.")

h2("3. About-Matt — close on a weak verb")
cur("If you want to see what I can do for your business, let's chat.")
rew("See a free preview of your site — no commitment.")
why("\"Let's chat\" is a soft, vague CTA the skill flags. End the most persuasive block on the page with the same "
    "concrete, zero-risk action as everywhere else.")
body("Also: the block is strong but runs long, and the tagline \"Agency results, a fraction of the price\" appears twice "
    "on the page. Keep it once (as a divider) and cut the duplicate — one idea per section.")

h2("4. Eyebrow / bundle line")
cur("Everything other agencies upsell you on - website, SEO, social, reviews - in one low monthly price.")
rew("Everything agencies sell you separately — website, SEO, social, reviews — in one monthly price from $99.")
why("Specificity ($99 beats \"low\") and cleaner em-dashes. \"Upsell you on\" is slightly clumsy; \"sell you separately\" "
    "lands the bundled-vs-piecemeal contrast more clearly.")

h2("5. Frank's testimonial — one consistency check")
flag("The card label reads \"Accordion player · Toronto,\" but Frank is positioned as a local (Niagara/his own site) "
     "performer. Make the location consistent with how he's referenced elsewhere so it doesn't read as a mismatch.")
why("Honest/consistent detail. The testimonial itself is excellent — before state, action, specific result, timeframe, "
    "emotion (\"completely painless\"). Don't touch the words; just reconcile the label.")

h2("6. Sub-page parity (Niagara, Affordable, One-Page)")
body("All three now carry the guarantee line in the CTA — good — but they're otherwise a full step behind the homepage: "
    "no testimonial, no founder block, no proof. A visitor landing on /web-design-niagara/ from search sees a weaker, "
    "less credible page than the homepage.")
bullets([
    "Add Frank's testimonial (or a one-line proof bar) to /web-design-niagara/ and /affordable-website-design/.",
    "Add a 2–3 line \"Built by Matt, here in Welland\" trust line with the Jetta Grove proof to those two pages.",
    "Move the guarantee into the hero on each sub-page, not just the bottom CTA.",
])
flag("One-Page page: confirm the hero H1. If it's \"When a one-page website works well,\" that's a weak, non-keyword H1 "
     "for \"one-page websites\" search intent — lead with \"One-page websites that make the next step obvious\" instead.")

# =====================================================================
h1("Part B — CRO pass")
sub("Skill: cro. Page type: homepage (serves cold search + warm referral). Primary goal: free-preview request (lead). "
    "Assumed traffic: local organic + direct. Analyzed in the framework's impact order.")

h2("Value proposition & headline (highest impact)")
body("Clear within 5 seconds: yes — \"get found and chosen,\" published prices, 24h, guarantee. Strong. The only lift is "
    "the headline-specificity tweak in Part A. Value prop clarity is no longer the bottleneck; trust depth and the path "
    "to action are.")

h2("Quick wins (do now)")
bullets([
    "Put the value in the button: \"Get my free preview\" (see Part A).",
    "Repeat the primary CTA at each decision point — after the comparison table, after Frank's testimonial, after About, and after pricing. Long pages lose people between CTAs.",
    "Dedupe the \"Agency results, a fraction of the price\" tagline.",
    "Add Google review stars + count near the hero as soon as you have 5–10 (you have a strong client in Frank — ask him first).",
    "Shorten the meta title (currently 77 chars, truncates in search) to under ~65.",
    "Add image alt text (flagged: no_image_alt) — accessibility + local image SEO.",
])

h2("High-impact changes (prioritize)")
bullets([
    "Add an instant booking option beside the form — \"Book a 15-min call\" (Cal.com). Right now even your highest-intent visitor has to wait for a reply; a competitor (Cool Koala) lets them book on the spot.",
    "Bring proof + guarantee onto the Niagara and Affordable sub-pages (Part A). These are your SEO landing pages; they convert worse than the homepage today.",
    "Add a short homepage FAQ to handle the live objections: \"What do you need from me?\", \"Is a 24-hour site any good?\", \"Can I cancel?\", \"Do I own it?\" The answers exist across the page — concentrating them removes hesitation before the CTA.",
    "Test moving the About block below \"How it works,\" so the scan order builds momentum: hero → proof table → testimonial → how it works → About (credibility) → pricing → CTA.",
])

h2("Trust & social proof")
bullets([
    "You've gone from zero proof to one strong, attributed testimonial + founder track record — the single biggest gap from the last audit is now half-closed.",
    "Next: review stars (none yet), and a second proof point as prototypes convert. A simple \"1 happy client, more on the way\" honesty beats fake volume — but get those Google reviews flowing.",
])

h2("Objection handling")
bullets([
    "Strong: guarantee, published pricing, comparison table, cancel-anytime, free preview.",
    "Gap: the \"is AI-built / 24-hour quality real?\" doubt isn't addressed head-on. The FAQ above is where to answer it — lean on the named mechanism and Frank's result.",
])

h2("Friction")
bullets([
    "Keep the lead form to 3–4 fields (name, business, current link, main goal). Confirm it isn't longer.",
    "Form-only is the main friction — add the booking option (above).",
    "Technical: title-too-long, no image alt, low content rate (improving — now ~847 words), and render-blocking CSS are all flagged. Confirm GA4/Hotjar actually fire before running any test below.",
])

h2("Test ideas (validate, don't assume)")
bullets([
    "Hero: current vs. timeframe-added vs. free-preview-led (Part A variants). Metric: preview requests.",
    "CTA button: \"Get my free preview\" vs. \"See my free website preview\" vs. current.",
    "Booking-first vs. form-first for high-intent traffic.",
    "About-block placement: high (current) vs. below How-it-works.",
    "Review-stars present vs. absent near hero (once you have them).",
])

# =====================================================================
h1("Part C — Priority order")
bullets([
    "1. Button copy → \"Get my free preview\" + repeat CTAs down the page (quick, high leverage).",
    "2. Add instant booking beside the form.",
    "3. Bring testimonial + founder proof + hero guarantee onto the Niagara and Affordable pages.",
    "4. Hero headline test (add the timeframe/specificity back).",
    "5. Add a short FAQ + Google review stars as they come in.",
    "6. Technical: meta title length, image alt, confirm analytics fire.",
])
flag("Open inputs: confirm the hero + final CTA button labels; confirm the One-Page hero H1; reconcile Frank's location "
     "label; confirm the lead form length; confirm GA4/Hotjar are firing.")

out = str(pathlib.Path(__file__).parent / "Well-and-Good-Websites-Copy-Pass-2-2026-06-29.docx")
doc.save(out)
print("Saved:", out)
