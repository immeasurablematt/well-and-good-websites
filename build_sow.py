#!/usr/bin/env python3
"""Build the per-tier Statement of Work (SOW) doc as a .docx."""
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11)
ACCENT = RGBColor(0x1A, 0x53, 0x4F); MUTED = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xB0, 0x00, 0x00); GREEN = RGBColor(0x1E, 0x7A, 0x33)

def title(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(24)
def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)
def body(t): return doc.add_paragraph(t)
def sub(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = True; r.font.color.rgb = MUTED; return p
def bullets(items):
    for it in items: doc.add_paragraph(it, style="List Bullet")
def numbered(items):
    for it in items: doc.add_paragraph(it, style="List Number")
def kv(lbl, value):
    p = doc.add_paragraph(); r = p.add_run(f"{lbl}: "); r.bold = True; r.font.color.rgb = ACCENT
    p.add_run(value); return p
def market(t):
    p = doc.add_paragraph(); r = p.add_run("vs. the market, "); r.bold = True; r.font.color.rgb = GREEN
    p.add_run(t); return p
def notincl(items):
    p = doc.add_paragraph(); r = p.add_run("Explicitly NOT included (add-on or higher tier):"); r.bold = True; r.font.color.rgb = RED
    for it in items: doc.add_paragraph(it, style="List Bullet")
def flag(t):
    p = doc.add_paragraph(); r = p.add_run(f"⚑ {t}"); r.bold = True; r.font.color.rgb = RED
def note(t):
    p = doc.add_paragraph(); r = p.add_run(f"→ {t}"); r.italic = True; r.font.color.rgb = MUTED
def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""; run = c.paragraphs[0].add_run(hd); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row): cells[i].text = val
    doc.add_paragraph("")
    return t

# =====================================================================
title("Well and Good Websites: Statement of Work by Tier")
sub("Scope-of-work specification for every pricing tier. Two product lines: (1) recurring growth bundles "
    "(Launch / Grow / Dominate), (2) one-time builds (Express / Starter / Standard / Premium) + optional Care. "
    "Written 2026-06-28. Designed to be both your delivery checklist and the basis of a client agreement.")

flag("ASSUMPTIONS TO CONFIRM, these specs are sensible, competitive defaults; the numbers below are your operational "
     "call. Adjust any before this goes client-facing: revision-round counts, content word counts (assumed 600–1000), "
     "citation counts (10 / 25 / 35+), tracked-keyword counts (5 / 15 / 30+), link-building targets (2–4/mo on Dominate), "
     "edit-turnaround SLAs (2-day / 1-day / priority), small-edit allowances, and every price on the Add-On Rate Card.")

# ---- shared terms ----
h1("0. Shared terms (apply to all tiers)")

h2("Definitions (so 'one-page' and 'content page' mean something)")
table(["Term", "Definition"], [
    ["One-page site", "A single scrolling page, up to 6 sections (hero, services, about, reviews, gallery, contact/map)."],
    ["Standard site", "Up to 5 linked pages (Home, Services, About, Reviews/Gallery, Contact) + one booking or quote flow."],
    ["Premium site", "Up to 8 linked pages + custom positioning/messaging + advanced sections (case study, before/after, menu/order)."],
    ["Content page/post", "One SEO-targeted page or article: 600–1000 words, 1 primary keyword, optimized title/meta/headings, internal + outbound links."],
    ["Citation", "A consistent Name-Address-Phone (NAP) listing on a directory (Bing Places, Yelp, YellowPages, Apple Maps, industry dirs)."],
    ["Revision round", "One consolidated set of change requests submitted together (not unlimited back-and-forth)."],
    ["Small edit", "A text/photo/hours/price swap on an existing section, ~30 min or less. Not a new section, page, or redesign."],
])

h2("Included in every recurring plan")
bullets([
    "Fast static hosting, SSL, daily backups, uptime monitoring.",
    "Mobile-responsive build, tap-to-call, click-to-map.",
    "Custom domain connection (client owns the domain).",
    "GA4 analytics setup and event tracking.",
    "Month-to-month. No lock-in, cancel anytime. Annual prepay = ~10 months (2 months free).",
    "Guarantee: live in 24 hours or your first month is free; revisions until you're happy.",
])

h2("Client responsibilities (the 24-hour clock depends on these)")
body("The 24-hour build clock starts when ALL of the following are received and the setup fee / first payment has cleared:")
bullets([
    "Business name, services, hours, service area, and preferred contact method.",
    "Logo (or a request to create one, add-on), brand colours if any.",
    "Photos (or approval to use stock/AI imagery), and access to existing Google reviews.",
    "Google Business Profile access, domain/DNS access, and any existing social accounts.",
])
note("Define the guarantee window precisely before publishing it: 24 hours = one business day from receiving all "
     "required assets + access. This keeps 'or your first month is free' safe to promise.")

h2("Acceptance & revisions")
bullets([
    "A preview is delivered within the tier's build window; the client submits one consolidated revision list per round.",
    "Revision rounds are capped per tier (below); additional rounds are billed at the Add-On rate.",
    "The site is deemed accepted on launch + 7 days, or when the first post-launch edit request is resolved.",
])

# =====================================================================
h1("1. Recurring growth bundles (at-a-glance)")
sub("The build inside each bundle mirrors a one-time tier: Launch≈Starter (one-page), Grow≈Standard (multi-page + "
    "booking), Dominate≈Premium (8 pages + custom copy). Everything else is the growth engine on top.")
table(["Deliverable", "Launch, $99/mo", "Grow, $399/mo", "Dominate, $899/mo"], [
    ["Setup / onboarding", "$497 (or $990/yr)", "$997 (or $3,990/yr)", "$1,497 (or $8,990/yr)"],
    ["Website", "One-page, up to 6 sections", "Standard, up to 5 pages + booking/quote", "Premium, up to 8 pages + custom copy"],
    ["Build revision rounds", "2", "3", "4 + iterative"],
    ["Tracked local keywords", "5", "15", "30+"],
    ["Citations (cumulative)", "10", "25", "35+ ongoing"],
    ["Content pieces / month", ", ", "2", "4"],
    ["Google Business Profile", "Optimize + 1 post/mo", "+ weekly posts", "+ full management"],
    ["Social posts / month", "4 (1 platform)", "~12 (2 platforms)", "20+ (3–4 platforms) + weekly video"],
    ["Review management", "Monitor", "Request flow + responses", "Full reputation strategy"],
    ["Reporting", "Monthly email", "Monthly PDF + summary", "Monthly call + live dashboard"],
    ["Small-edit SLA", "2 / mo, 2 business days", "Unlimited small, 1 business day", "Unlimited, priority same/next day"],
])

# ----- LAUNCH -----
h1("2. Launch · $99/mo + $497 setup")
kv("Who it's for", "A local owner who needs to exist on Google, look credible, and capture calls, without an agency.")
h3("Website")
bullets(["One-page site, up to 6 sections, mobile-responsive, tap-to-call, click-to-map.",
         "2 revision rounds. Preview within 24 business hours of assets; live same day on approval."])
h3("Local SEO foundation")
bullets(["Keyword research: up to 5 priority local keywords.",
         "On-page optimization: titles, meta descriptions, headings, image alt, LocalBusiness schema.",
         "10 core citations built + NAP consistency check.",
         "GA4 setup."])
h3("Google Business Profile")
bullets(["Claim/optimize: categories, services, hours, service area, description, photos.",
         "1 GBP post per month."])
h3("Social")
bullets(["4 posts/month on 1 platform (client's choice): graphic + caption, scheduled."])
h3("Care & reporting")
bullets(["Hosting, SSL, backups, monitoring included.",
         "Up to 2 small edits/month, 2-business-day turnaround.",
         "Monthly automated performance email: sessions, GBP views, calls, ranking snapshot."])
notincl(["Blog / content pages (Grow+)", "Booking or quote flow (Grow+)",
         "Multi-platform or video social (Grow / Dominate)", "Paid ad management", "Professional photo/video shoot"])
body("Acceptance: 2 revision rounds; accepted on launch + 7 days.")
market("Agencies bill $800–$2,500/mo for local SEO ALONE and $950–$3,000/mo for social. Launch bundles a custom site, "
       "SEO foundation, GBP, and social for $99/mo, published, no contract. KMK and most locals won't even show a price.")

# ----- GROW -----
h1("3. Grow · $399/mo + $997 setup (Most popular)")
kv("Who it's for", "An owner who wants the site to actively generate calls, bookings, and quotes, and to climb on Google.")
h3("Website")
bullets(["Everything in Launch, plus a Standard site (up to 5 pages) with a booking or quote-request flow "
         "(Cal.com/Calendly or a structured quote form).",
         "3 revision rounds."])
h3("Active local SEO")
bullets(["Expand to 15 tracked local keywords; monthly on-page improvements and internal linking.",
         "2 content pieces/month (600–900 words each), targeting local intent.",
         "Citations expanded to 25 total; quarterly NAP audit.",
         "Monthly competitor ranking monitoring."])
h3("Google Business Profile")
bullets(["Weekly GBP posts; Q&A seeding; photo cadence."])
h3("Social")
bullets(["~12 posts/month across 2 platforms: graphics + short copy, scheduled on a content calendar."])
h3("Review management")
bullets(["Automated review-request flow (post-job SMS/email).",
         "Monitor and respond to new reviews; monthly review summary."])
h3("Care & reporting")
bullets(["Unlimited small edits, 1-business-day turnaround.",
         "Monthly PDF report: traffic, tracked-keyword movement, GBP insights, leads/calls, social reach + a written summary."])
notincl(["Weekly video / short-form production (Dominate)", "Aggressive link building & outreach (Dominate)",
         "Monthly strategy call + live dashboard (Dominate)", "Paid ad management", "Custom positioning/messaging workshop (Dominate)"])
body("Acceptance: 3 revision rounds; accepted on launch + 7 days.")
market("This replaces an SEO retainer ($800–2,500) PLUS a social retainer ($950–3,000) PLUS maintenance, a bundle "
       "agencies would invoice at $1,800–$5,500/mo. Grow is $399/mo, cancel anytime. Cool Koala's floor is a $2,500 "
       "one-time project with no ongoing growth attached.")

# ----- DOMINATE -----
h1("4. Dominate · $899/mo + $1,497 setup")
kv("Who it's for", "An owner who wants to own the local market: a premium site plus an aggressive, managed growth system.")
h3("Website")
bullets(["Everything in Grow, plus a Premium site (up to 8 pages) with custom positioning & messaging (copywriting).",
         "Advanced sections: case study, before/after, or menu/order.",
         "4 revision rounds + ongoing iterative improvements (conversion-rate optimization)."])
h3("Aggressive SEO")
bullets(["30+ tracked keywords; 4 content pieces/month (mix of pages + posts, 800–1000 words).",
         "Link building / outreach: target 2–4 quality local or industry links per month.",
         "Quarterly technical SEO audit; ongoing citation growth (35+)."])
h3("Social + video")
bullets(["20+ posts/month across 3–4 platforms.",
         "1 short-form video per week, client provides raw clips, or I script + edit footage they supply.",
         "[Confirm video workflow: edit-only vs. you filming. On-site filming is an add-on.]"])
h3("Strategy & reporting")
bullets(["Monthly 45-minute strategy call.",
         "Live KPI dashboard (Looker Studio): rankings, traffic, leads, social, reviews.",
         "Full reputation strategy + review funnel."])
h3("Care")
bullets(["Unlimited edits, priority queue, same/next-business-day turnaround."])
notincl(["Paid ad SPEND (billed separately) and ongoing ad management (add-on % of spend)",
         "On-site professional videography/photography production (add-on)",
         "E-commerce / online-store build (add-on, quoted)",
         "Custom web-app or booking-system development beyond standard integrations"])
body("Acceptance: 4 revision rounds + iterative; accepted on launch + 7 days.")
market("Full-service agencies run $2,500–$80k+ on projects plus monthly retainers, behind 6–12 month contracts and a "
       "'request a quote' wall. Dominate is an integrated, published-price growth system at $899/mo with no lock-in.")

# =====================================================================
h1("5. One-time builds (no recurring fee)")
sub("For owners who want a website now and will handle their own growth. Add Care ($39/mo) for hosting + small edits. "
    "Upgrade into a recurring plan anytime; the setup fee paid converts toward onboarding.")
table(["Spec", "Express $297", "Starter $497", "Standard $997", "Premium $1,497"], [
    ["Scope", "1-section landing (hero, contact, tap-to-call, map)", "One-page, up to 6 sections", "Up to 5 pages + booking/quote flow", "Up to 8 pages + custom copywriting & positioning"],
    ["On-page SEO", "Basic (title/meta)", "Standard (titles, meta, schema)", "Standard + per-page", "Advanced + content structure"],
    ["GBP setup", ", ", "Basic", "Optimized", "Optimized"],
    ["Revision rounds", "1", "2", "3", "4"],
    ["Timeline", "24 hours", "24 hours", "2–3 business days", "3–5 business days"],
    ["Copywriting", "Template + your text", "Light edit of your text", "Section copy written", "Full custom copy & positioning"],
])
h3("Care add-on, $39/mo")
bullets(["Hosting, SSL, daily backups, uptime monitoring.",
         "Up to 2 small edits/month, 2-business-day turnaround.",
         "Upgrade to Launch/Grow/Dominate anytime."])
notincl(["Ongoing SEO, content, social, reviews, or reporting, those live in the recurring plans.",
         "Anything beyond the listed page/section count (Add-On rate card)."])
market("Cool Koala starts at a $2,500 project; Cedar+Mint $4,800+; KMK is quote-only. Your one-time builds top out at "
       "$1,497 with custom copy, and every price is published.")

# =====================================================================
h1("6. What the setup / onboarding fee buys")
sub("Justifies the upfront fee and scopes it, so it isn't confused with the monthly.")
table(["Fee", "Covers"], [
    ["$497 (Launch)", "Discovery, asset collection, one-page build, GBP claim/optimize, 10 citations, SEO foundation, GA4, launch."],
    ["$997 (Grow)", "All of the above + multi-page build, booking/quote flow, review-automation setup, 2 social accounts set up, content calendar."],
    ["$1,497 (Dominate)", "All of the above + positioning/messaging workshop, premium 8-page build, Looker Studio dashboard, link-building groundwork, video workflow setup."],
])

# =====================================================================
h1("7. Add-On Rate Card (the scope-creep guardrail)")
sub("Publish or keep internal, either way it gives every out-of-scope ask a price, so extra work is paid, not absorbed.")
flag("All prices below are SUGGESTED defaults, set your own before using.")
table(["Add-on", "Suggested price"], [
    ["Extra revision round", "$79"],
    ["Extra content page / post", "$99 each"],
    ["Extra social platform (ongoing)", "$79 / mo"],
    ["Rush, same-business-day change", "$49"],
    ["Extra website page (one-time build)", "$149"],
    ["New section on an existing one-page site", "$129"],
    ["Logo creation", "from $149"],
    ["Professional photo / video shoot (on-site)", "from $299, quoted"],
    ["Paid ads management", "15% of ad spend, min $199/mo (ad budget billed separately)"],
    ["E-commerce / online store", "Quoted"],
])

# =====================================================================
h1("8. Billing & terms (summary)")
bullets([
    "Setup/onboarding fee due upfront; the 24-hour clock starts once it clears and all assets/access are received.",
    "Monthly plans billed in advance, month-to-month, cancel anytime with 30 days' notice.",
    "Annual prepay = ~10 months (2 free): Launch $990, Grow $3,990, Dominate $8,990.",
    "Client owns the domain and content. On cancellation, the site can be exported / pointed elsewhere (no hostage).",
    "Ad spend is always the client's, billed separately and directly where possible.",
    "Guarantee: live in 24 hours (one business day from assets + access) or the first month is free; revisions until satisfied.",
])
note("Open inputs to lock before this is client-facing: confirm the flagged assumption numbers; confirm you can hit "
     "the 24-hour window reliably; decide whether the Add-On Rate Card is public or quote-on-request; confirm the "
     "Dominate video workflow (edit-only vs. you filming).")

out = str(pathlib.Path(__file__).parent / "Well-and-Good-Websites-SOW-by-Tier-2026-06-28.docx")
doc.save(out)
print("Saved:", out)
